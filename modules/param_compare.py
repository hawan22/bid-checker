"""
Module A: 技术参数响应表比对
对比招标文件中的技术参数要求与投标响应表，自动高亮偏差。
支持 Excel(.xlsx) 和 Word(.docx) 格式。
"""
import re
import difflib
import openpyxl
import pandas as pd
from io import BytesIO
from dataclasses import dataclass, field
from typing import List, Optional, Tuple


@dataclass
class ParamRow:
    index: str          # 序号
    name: str           # 参数名称
    requirement: str    # 招标要求
    response: str       # 投标响应
    result: str         # "PASS" | "WARN" | "FAIL" | "MISSING"
    color: str          # "green" | "yellow" | "red" | "gray"
    reason: str         # 说明


@dataclass
class ParamCompareResult:
    total: int
    passed: int
    warned: int
    failed: int
    missing: int
    rows: List[ParamRow] = field(default_factory=list)
    summary: str = ""


# ─── 数值比较逻辑 ──────────────────────────────────────────────────────────────

_NUM_RE = re.compile(r'([≥≤><＞＜>=!＝]{0,2})\s*([0-9]+(?:\.[0-9]+)?)\s*([a-zA-Z%℃㎡㎜㎝㎞㎏kKwW/]*)')

def _parse_numeric_req(text: str) -> Optional[Tuple[str, float, str]]:
    """从要求文本中提取（比较符, 数值, 单位）。如 '≥5mm' → ('>=', 5.0, 'mm')"""
    m = _NUM_RE.search(text)
    if not m:
        return None
    op = m.group(1).replace('＞', '>').replace('＜', '<').replace('＝', '=') or '='
    val = float(m.group(2))
    unit = m.group(3).strip()
    return op, val, unit


def _compare_numeric(req_op: str, req_val: float, resp_text: str) -> Tuple[str, str]:
    """返回 (result, reason)。"""
    m = _NUM_RE.search(resp_text)
    if not m:
        return 'WARN', '响应中未找到可比较数值，请人工核查'
    resp_val = float(m.group(2))

    checks = {
        '>=': resp_val >= req_val,
        '<=': resp_val <= req_val,
        '>': resp_val > req_val,
        '<': resp_val < req_val,
        '=': abs(resp_val - req_val) < 1e-9,
    }
    op = req_op if req_op in checks else '='
    ok = checks[op]
    sym = {'>=': '≥', '<=': '≤', '>': '>', '<': '<', '=': '='}.get(op, op)
    if ok:
        return 'PASS', f'响应值 {resp_val} 满足要求 {sym}{req_val}'
    else:
        return 'FAIL', f'响应值 {resp_val} 不满足要求 {sym}{req_val}'


# ─── 相似度匹配 ───────────────────────────────────────────────────────────────

def _best_match(name: str, candidates: List[str], threshold: float = 0.6) -> Optional[str]:
    """用 difflib 模糊匹配最佳候选参数名。"""
    if not candidates:
        return None
    name_clean = re.sub(r'[\s（()）【】\[\]]', '', name)
    cand_clean = [re.sub(r'[\s（()）【】\[\]]', '', c) for c in candidates]
    matches = difflib.get_close_matches(name_clean, cand_clean, n=1, cutoff=threshold)
    if matches:
        idx = cand_clean.index(matches[0])
        return candidates[idx]
    # 包含关系兜底
    for i, c in enumerate(cand_clean):
        if name_clean in c or c in name_clean:
            return candidates[i]
    return None


# ─── 文件解析 ─────────────────────────────────────────────────────────────────

def _read_excel_df(file_bytes: bytes) -> pd.DataFrame:
    return pd.read_excel(BytesIO(file_bytes), header=0, dtype=str).fillna('')


def _read_first_table_from_docx(file_bytes: bytes) -> Optional[pd.DataFrame]:
    """从 Word 文档中提取第一张表。"""
    from docx import Document
    doc = Document(BytesIO(file_bytes))
    if not doc.tables:
        return None
    tbl = doc.tables[0]
    data = [[cell.text.strip() for cell in row.cells] for row in tbl.rows]
    if not data:
        return None
    header = data[0]
    rows = data[1:]
    return pd.DataFrame(rows, columns=header).fillna('')


def _detect_columns(df: pd.DataFrame, role: str) -> Tuple[str, str]:
    """
    role = 'tender' → 返回 (参数名列名, 要求列名)
    role = 'bid'    → 返回 (参数名列名, 响应列名)
    """
    cols = list(df.columns)
    cols_lower = [str(c).replace(' ', '').lower() for c in cols]

    NAME_KEYS = ['参数名称', '参数', '名称', '项目', '指标']
    REQ_KEYS  = ['招标要求', '技术要求', '要求', '规格']
    RESP_KEYS = ['投标响应', '响应值', '投标值', '技术响应', '响应', '说明']

    def find(keys):
        for key in keys:
            for i, c in enumerate(cols_lower):
                if key.lower() in c:
                    return cols[i]
        return cols[0] if cols else ''

    name_col = find(NAME_KEYS)
    val_col  = find(REQ_KEYS if role == 'tender' else RESP_KEYS)
    return name_col, val_col


# ─── 主函数 ───────────────────────────────────────────────────────────────────

def compare_params(
    tender_bytes: bytes,
    tender_name: str,
    bid_bytes: bytes,
    bid_name: str,
    tender_name_col: str = "",
    tender_req_col: str = "",
    bid_name_col: str = "",
    bid_resp_col: str = "",
) -> ParamCompareResult:
    """
    对比招标参数表与投标响应表。
    如果未指定列名，自动猜测。
    """
    # 解析文件
    def load_df(b, fname):
        if fname.endswith('.xlsx') or fname.endswith('.xls'):
            return _read_excel_df(b)
        elif fname.endswith('.docx'):
            return _read_first_table_from_docx(b)
        return None

    tender_df = load_df(tender_bytes, tender_name)
    bid_df    = load_df(bid_bytes, bid_name)

    if tender_df is None or bid_df is None:
        return ParamCompareResult(0, 0, 0, 0, 0, [],
            "文件解析失败，请确认为 .xlsx 或 .docx 格式")

    # 自动检测列
    if not tender_name_col or not tender_req_col:
        tender_name_col, tender_req_col = _detect_columns(tender_df, 'tender')
    if not bid_name_col or not bid_resp_col:
        bid_name_col, bid_resp_col = _detect_columns(bid_df, 'bid')

    tender_params = list(tender_df[tender_name_col].astype(str))
    bid_map = {
        str(row[bid_name_col]): str(row[bid_resp_col])
        for _, row in bid_df.iterrows()
    }

    rows = []
    for i, row in tender_df.iterrows():
        name = str(row[tender_name_col]).strip()
        req  = str(row[tender_req_col]).strip()
        if not name or name in ('nan', '参数名称', '序号'):
            continue

        # 在投标表中查找对应参数
        resp = bid_map.get(name)
        if resp is None:
            matched = _best_match(name, list(bid_map.keys()))
            resp = bid_map.get(matched, '') if matched else ''
            if not resp:
                rows.append(ParamRow(str(i+1), name, req, '', 'MISSING', 'gray', '投标响应表中未找到此参数'))
                continue

        # 判断结果
        numeric = _parse_numeric_req(req)
        if numeric:
            op, val, unit = numeric
            result, reason = _compare_numeric(op if op else '>=', val, resp)
        elif not resp.strip():
            result, reason = 'WARN', '投标响应为空，需补充'
        elif len(resp.strip()) < 3:
            result, reason = 'WARN', '响应内容过于简短，建议补充说明'
        else:
            # 检查是否存在明显的"未响应"关键词
            if any(kw in resp for kw in ['详见', '参见', '如图', '以实物为准']):
                result, reason = 'WARN', '响应为参引性说明，建议补充具体参数值'
            else:
                result, reason = 'PASS', '已响应'

        color_map = {'PASS': 'green', 'WARN': 'yellow', 'FAIL': 'red', 'MISSING': 'gray'}
        rows.append(ParamRow(str(i+1), name, req, resp, result, color_map[result], reason))

    passed  = sum(1 for r in rows if r.result == 'PASS')
    warned  = sum(1 for r in rows if r.result == 'WARN')
    failed  = sum(1 for r in rows if r.result == 'FAIL')
    missing = sum(1 for r in rows if r.result == 'MISSING')
    total   = len(rows)

    issues = []
    if failed:  issues.append(f'❌ {failed} 项不满足要求')
    if missing: issues.append(f'⬜ {missing} 项未响应')
    if warned:  issues.append(f'⚠️ {warned} 项需人工核查')
    if not issues:
        summary = f'✅ 全部 {total} 项参数均已响应，未发现明显问题。'
    else:
        summary = f'共 {total} 项，通过 {passed} 项；' + '；'.join(issues)

    return ParamCompareResult(total, passed, warned, failed, missing, rows, summary)
