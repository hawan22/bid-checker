"""
页眉页脚一致性检测模块
① Word: 提取各节页眉/页脚，检查项目名/公司名是否与正文一致
② Word: 检测不同节的页眉是否出现不一致（复制自其他标书）
③ PDF: 通过检测每页首行/末行的重复文本推断页眉/页脚
"""
import re
from dataclasses import dataclass, field
from typing import List, Dict, Optional, Set

try:
    from docx import Document
    DOCX_OK = True
except ImportError:
    DOCX_OK = False

try:
    import fitz
    FITZ_OK = True
except ImportError:
    FITZ_OK = False


# ─── 数据结构 ──────────────────────────────────────────────────

@dataclass
class HeaderFooterEntry:
    location: str     # "第N节页眉" / "第N节页脚" / "PDF推断页眉"
    content: str
    section_num: int


@dataclass
class HFIssue:
    issue_type: str   # 不一致/与正文不符/疑似残留/页码异常
    description: str
    locations: List[str]
    risk_level: str   # 高/中/低
    suggestion: str


@dataclass
class HeaderFooterResult:
    headers: List[HeaderFooterEntry] = field(default_factory=list)
    footers: List[HeaderFooterEntry] = field(default_factory=list)
    issues: List[HFIssue] = field(default_factory=list)
    problems: List[str] = field(default_factory=list)
    summary: str = ""
    file_type: str = ""
    score: int = 100


# ─── Word 页眉/页脚提取 ────────────────────────────────────────

def _get_word_headers_footers(file_bytes: bytes):
    """提取 Word 所有节的页眉和页脚文本"""
    if not DOCX_OK:
        return [], []
    import io
    headers, footers = [], []
    try:
        doc = Document(io.BytesIO(file_bytes))
        for i, section in enumerate(doc.sections):
            # 页眉
            hdr = section.header
            if hdr and not hdr.is_linked_to_previous:
                text = "\n".join(p.text.strip() for p in hdr.paragraphs if p.text.strip())
                if text:
                    headers.append(HeaderFooterEntry(
                        location=f"第{i+1}节页眉", content=text, section_num=i+1))
            # 页脚
            ftr = section.footer
            if ftr and not ftr.is_linked_to_previous:
                text = "\n".join(p.text.strip() for p in ftr.paragraphs if p.text.strip())
                if text:
                    footers.append(HeaderFooterEntry(
                        location=f"第{i+1}节页脚", content=text, section_num=i+1))
    except Exception:
        pass
    return headers, footers


# ─── PDF 页眉/页脚推断 ────────────────────────────────────────

def _infer_pdf_headers_footers(pages: List[str]) -> tuple:
    """通过多页重复出现的首行/末行推断页眉/页脚"""
    if not pages or len(pages) < 3:
        return [], []

    # 提取每页首行和末行
    first_lines, last_lines = [], []
    for pg in pages:
        lines = [l.strip() for l in pg.split('\n') if l.strip()]
        if lines:
            first_lines.append(lines[0])
            last_lines.append(lines[-1])

    headers, footers = [], []

    # 首行重复出现 ≥ 页数的50% → 疑似页眉
    from collections import Counter
    fl_counter = Counter(first_lines)
    for text, count in fl_counter.most_common(2):
        if count >= max(2, len(pages) * 0.4) and len(text) > 2:
            headers.append(HeaderFooterEntry(
                location=f"PDF推断页眉（{count}/{len(pages)}页重复）",
                content=text, section_num=0))

    ll_counter = Counter(last_lines)
    for text, count in ll_counter.most_common(2):
        if count >= max(2, len(pages) * 0.4) and len(text) > 2:
            # 排除纯页码
            if not re.match(r'^\d+$', text):
                footers.append(HeaderFooterEntry(
                    location=f"PDF推断页脚（{count}/{len(pages)}页重复）",
                    content=text, section_num=0))

    return headers, footers


# ─── 一致性分析 ───────────────────────────────────────────────

def _extract_company_name(text: str) -> str:
    """从正文提取公司名（最常见的有限公司名）"""
    from collections import Counter
    pattern = re.compile(
        r'([\u4e00-\u9fa5a-zA-Z0-9（）]{2,20}'
        r'(?:有限公司|股份有限公司|集团有限公司|有限责任公司))'
    )
    matches = pattern.findall(text[:5000])  # 只看前5000字符
    if not matches:
        return ""
    counter = Counter(matches)
    return counter.most_common(1)[0][0]


def _extract_project_name(text: str) -> str:
    """从正文提取项目名"""
    m = re.search(r'项目名称[：:：\s]*([\u4e00-\u9fa5a-zA-Z0-9（）]{4,40})', text)
    return m.group(1).strip() if m else ""


def _analyze_consistency(
    headers: List[HeaderFooterEntry],
    footers: List[HeaderFooterEntry],
    text: str,
) -> List[HFIssue]:
    issues: List[HFIssue] = []

    all_hf = headers + footers
    if not all_hf:
        return issues

    body_company = _extract_company_name(text)
    body_project = _extract_project_name(text)

    # ── 1. 页眉之间不一致（多节不同）────────────────────────────
    header_contents = [h.content for h in headers]
    if len(set(header_contents)) > 1:
        # 去除页码差异，检查文字部分
        unique = list(set(header_contents))
        issues.append(HFIssue(
            issue_type="页眉内容不一致",
            description=f"不同节的页眉内容不同，共 {len(unique)} 种：\n" +
                        "\n".join(f"  • 「{c[:40]}」" for c in unique[:4]),
            locations=[h.location for h in headers],
            risk_level="中",
            suggestion="统一所有节的页眉内容；若是刻意分节（封面无页眉），可忽略首节",
        ))

    footer_contents = [f.content for f in footers]
    if len(set(footer_contents)) > 1:
        unique = list(set(footer_contents))
        issues.append(HFIssue(
            issue_type="页脚内容不一致",
            description=f"不同节的页脚内容不同，共 {len(unique)} 种：\n" +
                        "\n".join(f"  • 「{c[:40]}」" for c in unique[:4]),
            locations=[f.location for f in footers],
            risk_level="中",
            suggestion="统一所有节的页脚格式",
        ))

    # ── 2. 页眉/脚中的公司名与正文不符 ─────────────────────────
    if body_company:
        for hf in all_hf:
            # 找出页眉/脚中出现的公司名
            hf_company_m = re.search(
                r'([\u4e00-\u9fa5a-zA-Z0-9（）]{2,20}'
                r'(?:有限公司|股份有限公司|集团有限公司|有限责任公司))',
                hf.content
            )
            if hf_company_m:
                hf_company = hf_company_m.group(1)
                if hf_company != body_company:
                    issues.append(HFIssue(
                        issue_type="页眉/脚公司名与正文不符",
                        description=(
                            f"{hf.location} 中公司名为「{hf_company}」，"
                            f"与正文中「{body_company}」不一致\n"
                            "⚠️ 可能是从其他公司的标书模板复制而来未修改页眉"
                        ),
                        locations=[hf.location],
                        risk_level="高",
                        suggestion=f"将{hf.location}中的公司名改为「{body_company}」",
                    ))

    # ── 3. 页眉/脚中含项目名，与正文项目名不符 ──────────────────
    if body_project:
        for hf in all_hf:
            # 检查是否有其他项目名关键词
            project_kw = re.search(r'[\u4e00-\u9fa5]{3,15}(?:项目|工程|采购|招标)', hf.content)
            if project_kw:
                hf_project = project_kw.group(0)
                # 简单相似度：核心词是否有交集
                body_words = set(re.findall(r'[\u4e00-\u9fa5]{2,4}', body_project))
                hf_words = set(re.findall(r'[\u4e00-\u9fa5]{2,4}', hf_project))
                overlap = body_words & hf_words
                if len(overlap) < 1 and len(body_words) >= 2:
                    issues.append(HFIssue(
                        issue_type="页眉/脚项目名与正文不符",
                        description=(
                            f"{hf.location} 中项目名「{hf_project}」"
                            f"与正文「{body_project}」差异较大"
                        ),
                        locations=[hf.location],
                        risk_level="中",
                        suggestion=f"核查并统一{hf.location}中的项目名称",
                    ))

    # ── 4. 页眉/脚疑似包含其他竞争对手/模板残留关键词 ────────────
    suspicious_keywords = ["模板", "示例", "样本", "XXX", "（请填", "【", "TODO"]
    for hf in all_hf:
        for kw in suspicious_keywords:
            if kw in hf.content:
                issues.append(HFIssue(
                    issue_type="页眉/脚含模板残留",
                    description=f"{hf.location} 含疑似模板标记：「{kw}」，内容：「{hf.content[:40]}」",
                    locations=[hf.location],
                    risk_level="高",
                    suggestion="删除或替换页眉/脚中的模板标记",
                ))
                break

    return issues


# ─── 主入口 ───────────────────────────────────────────────────

def check_header_footer(
    file_bytes: bytes,
    filename: str,
    text: str,
    pages: List[str] = None,
) -> HeaderFooterResult:
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    result = HeaderFooterResult(file_type=ext)

    if ext in ("docx", "doc") and file_bytes and DOCX_OK:
        result.headers, result.footers = _get_word_headers_footers(file_bytes)
    elif ext == "pdf" and pages:
        result.headers, result.footers = _infer_pdf_headers_footers(pages)

    all_hf = result.headers + result.footers

    if all_hf and text:
        result.issues = _analyze_consistency(result.headers, result.footers, text)

    # 构建 problems & 评分
    score = 100
    for issue in result.issues:
        result.problems.append(
            f"{'🔴' if issue.risk_level=='高' else '🟠'} {issue.issue_type}：{issue.description[:60]}"
        )
        score -= 25 if issue.risk_level == "高" else 12
    result.score = max(0, score)

    n_hdr = len(result.headers)
    n_ftr = len(result.footers)
    n_issues = len(result.issues)

    if not all_hf:
        if ext in ("docx", "doc"):
            result.summary = "ℹ️ 文档未设置页眉页脚（或均为空白）"
        else:
            result.summary = "ℹ️ PDF未检测到重复页眉/页脚内容"
    elif n_issues == 0:
        result.summary = f"✅ 页眉/页脚一致（{n_hdr}处页眉、{n_ftr}处页脚），与正文内容吻合"
    else:
        high = sum(1 for i in result.issues if i.risk_level == "高")
        mid  = sum(1 for i in result.issues if i.risk_level == "中")
        parts = []
        if high: parts.append(f"🔴 高风险 {high} 项")
        if mid:  parts.append(f"🟠 中风险 {mid} 项")
        result.summary = "发现页眉/页脚问题：" + "、".join(parts)

    return result
