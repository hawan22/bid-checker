"""
OCR Helper
支持从图片和 PDF 扫描页提取文字，使用本地 Tesseract（中文+英文）。
调用前先检查 Tesseract 是否可用，不可用时给出友好提示。
"""
import io
import os
import re
from pathlib import Path
from typing import Optional

# Tesseract 路径（Windows 默认安装位置）
TESSERACT_PATHS = [
    r"C:\Program Files\Tesseract-OCR\tesseract.exe",
    r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
    "tesseract",  # 已在 PATH 中
]

_tesseract_ok: Optional[bool] = None
_tesseract_path: Optional[str] = None


def _find_tesseract() -> Optional[str]:
    global _tesseract_ok, _tesseract_path
    if _tesseract_ok is not None:
        return _tesseract_path if _tesseract_ok else None

    import shutil
    for p in TESSERACT_PATHS:
        if os.path.isfile(p) or shutil.which(p):
            _tesseract_ok = True
            _tesseract_path = p
            return p

    _tesseract_ok = False
    return None


def tesseract_available() -> bool:
    return _find_tesseract() is not None


def deskew_image(img) -> tuple:
    """
    自动检测并纠正图片偏斜（手机拍照倾斜）。
    使用 Tesseract OSD（方向与脚本检测）。
    返回 (纠正后的PIL Image, 旋转角度)；失败返回原图+0。
    """
    tess_path = _find_tesseract()
    if not tess_path:
        return img, 0

    try:
        import pytesseract
        pytesseract.pytesseract.tesseract_cmd = tess_path
        # OSD 检测需要一定文字量，小图跳过
        if img.width < 200 or img.height < 200:
            return img, 0
        osd = pytesseract.image_to_osd(img, config="--psm 0 -c min_characters_to_try=5")
        angle_match = re.search(r'Rotate:\s*(\d+)', osd)
        if angle_match:
            angle = int(angle_match.group(1))
            if angle in (90, 180, 270):
                img = img.rotate(-angle, expand=True)
                return img, angle
    except Exception:
        pass
    return img, 0


def ocr_image_bytes(image_bytes: bytes, lang: str = "chi_sim+eng",
                    auto_deskew: bool = True) -> str:
    """
    对图片字节流进行 OCR，返回识别文本。
    auto_deskew=True 时自动纠正手机拍照旋转偏斜。
    lang: Tesseract 语言参数，默认中文+英文
    """
    tess_path = _find_tesseract()
    if not tess_path:
        return ""

    import pytesseract
    pytesseract.pytesseract.tesseract_cmd = tess_path

    from PIL import Image
    img = Image.open(io.BytesIO(image_bytes))
    # 预处理：灰度 + 适当放大（提高识别率）
    img = img.convert("L")
    if max(img.size) < 1000:
        scale = 1500 / max(img.size)
        new_size = (int(img.width * scale), int(img.height * scale))
        img = img.resize(new_size, Image.LANCZOS)

    # 自动纠偏（手机拍照90/180/270度旋转）
    if auto_deskew:
        img, _ = deskew_image(img)

    text = pytesseract.image_to_string(img, lang=lang, config="--psm 6")
    return text.strip()


def ocr_pdf_page(page) -> str:
    """对 PyMuPDF 的一个页面对象进行 OCR（适用于图片型页面）。"""
    tess_path = _find_tesseract()
    if not tess_path:
        return ""

    # 将页面渲染为图片（分辨率 200 DPI）
    mat = __import__("fitz").Matrix(200 / 72, 200 / 72)
    pix = page.get_pixmap(matrix=mat, colorspace=__import__("fitz").csGRAY)
    img_bytes = pix.tobytes("png")
    return ocr_image_bytes(img_bytes)


def extract_text_from_image(file_bytes: bytes, filename: str = "") -> tuple[str, str]:
    """
    从图片文件提取文字。
    返回 (text, method) 其中 method = "ocr" | "error"
    """
    if not tesseract_available():
        return "", "error"
    text = ocr_image_bytes(file_bytes)
    return text, "ocr"


def extract_text_from_docx(file_bytes: bytes) -> str:
    """从 Word(.docx) 文件提取全部文字。"""
    from docx import Document
    doc = Document(io.BytesIO(file_bytes))
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)
    return "\n".join(parts)


def extract_text_smart(file_bytes: bytes, filename: str) -> tuple[str, str, list[dict]]:
    """
    智能提取文字，支持 PDF/图片/Word。
    返回 (全文, 方法说明, 逐页/逐段列表)
    逐页列表: [{"page": N, "text": "...", "method": "text"|"ocr"}]
    """
    ext = Path(filename).suffix.lower()
    pages = []

    if ext in (".jpg", ".jpeg", ".png", ".bmp", ".tiff", ".tif", ".webp"):
        text, method = extract_text_from_image(file_bytes, filename)
        if method == "error":
            return "", "❌ Tesseract 未安装，无法识别图片", []
        pages = [{"page": 1, "text": text, "method": "ocr",
                  "note": f"OCR识别，共 {len(text)} 字符"}]
        return text, "OCR（图片识别）", pages

    elif ext == ".docx":
        text = extract_text_from_docx(file_bytes)
        pages = [{"page": 1, "text": text, "method": "text",
                  "note": f"Word文档，共 {len(text)} 字符"}]
        return text, "Word文档直读", pages

    elif ext == ".pdf":
        import fitz
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        full_text = ""
        for i, page in enumerate(doc):
            t = page.get_text("text").strip()
            if len(t) >= 30:
                method = "text"
                note = f"文字层，{len(t)} 字符"
            else:
                # 图片页，尝试 OCR
                if tesseract_available():
                    t = ocr_pdf_page(page)
                    method = "ocr"
                    note = f"图片页OCR，{len(t)} 字符"
                else:
                    method = "image_no_ocr"
                    note = "图片页，Tesseract 未安装，跳过OCR"
            pages.append({"page": i + 1, "text": t, "method": method, "note": note})
            full_text += t + "\n"
        doc.close()
        has_ocr = any(p["method"] == "ocr" for p in pages)
        method_desc = "PDF（文字层+OCR）" if has_ocr else "PDF文字层"
        return full_text, method_desc, pages

    return "", f"不支持的文件格式：{ext}", []
